#!/usr/bin/env python3
"""Convert project Markdown documents into styled Word documents.

This is intentionally small and deterministic for the project's PRD/analysis docs:
headings, paragraphs, tables, lists, block quotes, fenced code blocks, and inline
bold/code formatting are supported without relying on external markdown packages.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ACCENT = "2563EB"
ACCENT_DARK = "1E3A8A"
TEXT = "1F2937"
MUTED = "6B7280"
BORDER = "D6DEE8"
TABLE_HEADER = "EAF2FF"
CALLOUT = "F4F8FF"
CODE_BG = "F6F8FA"
BODY_FONT = "PingFang SC"
MONO_FONT = "Menlo"


def apply_run_font(run, font_name: str = BODY_FONT) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Heading 1", 18, ACCENT_DARK, 16, 8),
        ("Heading 2", 15, ACCENT_DARK, 14, 6),
        ("Heading 3", 12.5, TEXT, 10, 4),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_inline_markdown(paragraph, text: str, *, color: str | None = None) -> None:
    """Add inline bold/code runs while preserving plain text."""
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            apply_run_font(run)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_run_font(run)
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            apply_run_font(run, MONO_FONT)
            run.font.size = Pt(9.2)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        apply_run_font(run)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def split_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|").split("|")
    return [cell.strip() for cell in raw]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = cells[c_idx]
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            add_inline_markdown(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(9.2)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F9FBFD")
    doc.add_paragraph()


def add_cover(doc: Document, title: str, subtitle_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    apply_run_font(run)

    for line in subtitle_lines:
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        add_inline_markdown(meta, line.replace("  ", ""), color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_shading(rule, ACCENT)
    doc.add_section(WD_SECTION_START.CONTINUOUS)


def convert_markdown(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = md_path.stem
    subtitle: list[str] = []
    idx = 0
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        idx = 1
        while idx < len(lines) and (not lines[idx].startswith("## ")):
            if lines[idx].strip():
                subtitle.append(lines[idx].strip())
            idx += 1
    add_cover(doc, title, subtitle)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Cm(0.25)
            p.paragraph_format.right_indent = Cm(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, CODE_BG)
            run = p.add_run("\n".join(code_lines))
            apply_run_font(run, MONO_FONT)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(TEXT)
            code_lines = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
                code_lines = []
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            next_line = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
            if is_table_separator(stripped):
                idx += 1
                continue
            table_rows.append(split_table_row(stripped))
            idx += 1
            continue
        else:
            flush_table()

        if not stripped:
            idx += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            doc.add_paragraph(heading_match.group(2), style=f"Heading {level}")
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, CALLOUT)
            add_inline_markdown(p, stripped.lstrip("> ").strip(), color=TEXT)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, stripped)
        idx += 1

    flush_table()
    flush_code()

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "灵栖胶囊 Capsule"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8.5)
        footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: export_markdown_to_docx.py input.md output.docx", file=sys.stderr)
        return 2
    convert_markdown(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
